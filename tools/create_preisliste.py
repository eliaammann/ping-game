from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "documents"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "Preisliste_Ping-Spiel.docx"
PDF_PATH = OUT / "Preisliste_Ping-Spiel.pdf"

NAVY = "17324D"
BLUE = "2B6F93"
PALE_BLUE = "EAF4F8"
PALE_GOLD = "FFF4D6"
GOLD = "E8AA2C"
INK = RGBColor(23, 50, 77)
MUTED = RGBColor(83, 99, 112)
WHITE = RGBColor(255, 255, 255)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color="D4E0E6", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = tc_borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_cell_margins(cell, top=130, start=170, bottom=130, end=170):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_text(paragraph, text, size, bold=False, color=INK, font="Aptos"):
    run = paragraph.add_run(text)
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    return run


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.55)
section.bottom_margin = Cm(1.45)
section.left_margin = Cm(1.65)
section.right_margin = Cm(1.65)
section.header_distance = Cm(0.7)
section.footer_distance = Cm(0.7)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.paragraph_format.space_after = Pt(0)
add_text(hp, "PING-SPIEL", 8.5, True, MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(0)
add_text(fp, "Alle Preise gelten pro Kauf.", 8.5, False, MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
add_text(p, "PREISLISTE", 28, True, INK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
add_text(p, "Credits, Zonenverkleinerung & Hinweise", 12, False, MUTED)

credit = doc.add_table(rows=1, cols=2)
credit.alignment = WD_TABLE_ALIGNMENT.CENTER
credit.autofit = False
credit.columns[0].width = Cm(11.8)
credit.columns[1].width = Cm(4.9)
left, right = credit.rows[0].cells
left.width, right.width = Cm(11.8), Cm(4.9)
for cell in (left, right):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=210, bottom=210)
    borders(cell, NAVY, "12")
shade(left, NAVY)
shade(right, GOLD)
lp = left.paragraphs[0]
lp.paragraph_format.space_after = Pt(0)
add_text(lp, "1 BÄNDEL", 18, True, WHITE)
rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp.paragraph_format.space_after = Pt(0)
add_text(rp, "= 1 CREDIT", 17, True, INK)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(17)
p.paragraph_format.space_after = Pt(7)
add_text(p, "ZONENVERKLEINERUNG", 14, True, INK)

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.columns[0].width = Cm(12.2)
table.columns[1].width = Cm(4.5)
headers = table.rows[0].cells
set_repeat_table_header(table.rows[0])
for cell, text, align in zip(headers, ("Stufe", "Preis"), (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER)):
    shade(cell, BLUE)
    borders(cell, BLUE, "10")
    set_cell_margins(cell, top=125, bottom=125)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    par = cell.paragraphs[0]
    par.alignment = align
    par.paragraph_format.space_after = Pt(0)
    add_text(par, text, 11, True, WHITE)

prices = [
    ("1. Stufe", "3 Credits"),
    ("2. Stufe", "5 Credits"),
    ("3. Stufe", "7 Credits"),
    ("4. Stufe", "15 Credits"),
]
for idx, (level, price) in enumerate(prices):
    cells = table.add_row().cells
    fill = "FFFFFF" if idx % 2 == 0 else PALE_BLUE
    for cell in cells:
        shade(cell, fill)
        borders(cell)
        set_cell_margins(cell, top=150, bottom=150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
    add_text(cells[0].paragraphs[0], level, 12, True, INK)
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
    add_text(cells[1].paragraphs[0], price, 12, True, INK)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(17)
p.paragraph_format.space_after = Pt(7)
add_text(p, "VERSTECK-HINWEISE", 14, True, INK)

extras = doc.add_table(rows=2, cols=2)
extras.alignment = WD_TABLE_ALIGNMENT.CENTER
extras.autofit = False
extras.columns[0].width = Cm(12.2)
extras.columns[1].width = Cm(4.5)
for row, (item, price) in zip(
    extras.rows,
    (("Schriftliche Versteckbeschreibung", "17 Credits"), ("Bild des Ortes", "20 Credits")),
):
    for cell in row.cells:
        shade(cell, "FFFFFF")
        borders(cell)
        set_cell_margins(cell, top=165, bottom=165)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
    add_text(row.cells[0].paragraphs[0], item, 11.5, True, INK)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
    add_text(row.cells[1].paragraphs[0], price, 11.5, True, INK)

spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(4)
spacer.paragraph_format.space_after = Pt(4)

note = doc.add_table(rows=1, cols=1)
note.alignment = WD_TABLE_ALIGNMENT.CENTER
note.autofit = False
note.columns[0].width = Cm(16.7)
cell = note.cell(0, 0)
shade(cell, PALE_GOLD)
borders(cell, GOLD, "14")
set_cell_margins(cell, top=190, bottom=190, start=220, end=220)
cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
add_text(p, "WICHTIG: Es kann keine Zonenstufe übersprungen werden.", 12.5, True, INK)

doc.core_properties.title = "Preisliste Ping-Spiel"
doc.core_properties.subject = "Credits und Preise"
doc.core_properties.author = "Ping-Spiel"
doc.core_properties.keywords = "Preisliste, Credits, Zonenverkleinerung"

doc.save(DOCX_PATH)


def pdf_text(c, x, y, text, size, bold=False, color=NAVY, align="left"):
    font = "Helvetica-Bold" if bold else "Helvetica"
    c.setFont(font, size)
    c.setFillColor(HexColor(f"#{color.lstrip('#')}"))
    if align == "center":
        c.drawCentredString(x, y, text)
    elif align == "right":
        c.drawRightString(x, y, text)
    else:
        c.drawString(x, y, text)


def rounded_box(c, x, y, w, h, fill, stroke="D4E0E6", radius=3 * mm, line=1):
    c.setFillColor(HexColor(f"#{fill.lstrip('#')}"))
    c.setStrokeColor(HexColor(f"#{stroke.lstrip('#')}"))
    c.setLineWidth(line)
    c.roundRect(x, y, w, h, radius, fill=1, stroke=1)


def hc(value):
    return HexColor(f"#{value.lstrip('#')}")


c = canvas.Canvas(str(PDF_PATH), pagesize=A4)
page_w, page_h = A4
margin = 17 * mm
content_w = page_w - 2 * margin

pdf_text(c, page_w - margin, page_h - 12 * mm, "PING-SPIEL", 8.5, True, "536370", "right")
pdf_text(c, page_w / 2, page_h - 31 * mm, "PREISLISTE", 28, True, NAVY, "center")
pdf_text(c, page_w / 2, page_h - 39 * mm, "Credits, Zonenverkleinerung & Hinweise", 12, False, "536370", "center")

y = page_h - 60 * mm
h = 18 * mm
left_w = content_w * 0.70
rounded_box(c, margin, y, left_w, h, NAVY, NAVY, 3 * mm, 1.2)
rounded_box(c, margin + left_w, y, content_w - left_w, h, GOLD, GOLD, 3 * mm, 1.2)
pdf_text(c, margin + 6 * mm, y + 6.3 * mm, "1 BÄNDEL", 18, True, "FFFFFF")
pdf_text(c, margin + left_w + (content_w - left_w) / 2, y + 6.4 * mm, "= 1 CREDIT", 17, True, NAVY, "center")

y -= 17 * mm
pdf_text(c, margin, y, "ZONENVERKLEINERUNG", 14, True, NAVY)
y -= 8 * mm
row_h = 12 * mm
col1 = content_w * 0.73
c.setFillColor(hc(BLUE))
c.rect(margin, y - row_h, content_w, row_h, fill=1, stroke=0)
pdf_text(c, margin + 5 * mm, y - 7.6 * mm, "Stufe", 11, True, "FFFFFF")
pdf_text(c, margin + col1 + (content_w - col1) / 2, y - 7.6 * mm, "Preis", 11, True, "FFFFFF", "center")
y -= row_h
for idx, (level, price) in enumerate(prices):
    fill = "FFFFFF" if idx % 2 == 0 else PALE_BLUE
    c.setFillColor(hc(fill))
    c.setStrokeColor(hc("D4E0E6"))
    c.setLineWidth(0.8)
    c.rect(margin, y - row_h, col1, row_h, fill=1, stroke=1)
    c.rect(margin + col1, y - row_h, content_w - col1, row_h, fill=1, stroke=1)
    pdf_text(c, margin + 5 * mm, y - 7.7 * mm, level, 12, True, NAVY)
    pdf_text(c, margin + col1 + (content_w - col1) / 2, y - 7.7 * mm, price, 12, True, NAVY, "center")
    y -= row_h

y -= 12 * mm
pdf_text(c, margin, y, "VERSTECK-HINWEISE", 14, True, NAVY)
y -= 8 * mm
for item, price in (("Schriftliche Versteckbeschreibung", "17 Credits"), ("Bild des Ortes", "20 Credits")):
    c.setFillColor(white)
    c.setStrokeColor(hc("D4E0E6"))
    c.rect(margin, y - row_h, col1, row_h, fill=1, stroke=1)
    c.rect(margin + col1, y - row_h, content_w - col1, row_h, fill=1, stroke=1)
    pdf_text(c, margin + 5 * mm, y - 7.7 * mm, item, 11.5, True, NAVY)
    pdf_text(c, margin + col1 + (content_w - col1) / 2, y - 7.7 * mm, price, 11.5, True, NAVY, "center")
    y -= row_h

y -= 12 * mm
note_h = 17 * mm
rounded_box(c, margin, y - note_h, content_w, note_h, PALE_GOLD, GOLD, 3 * mm, 1.4)
pdf_text(
    c,
    page_w / 2,
    y - 10.4 * mm,
    "WICHTIG: Es kann keine Zonenstufe übersprungen werden.",
    12.5,
    True,
    NAVY,
    "center",
)

pdf_text(c, page_w / 2, 10 * mm, "Alle Preise gelten pro Kauf.", 8.5, False, "536370", "center")
c.setTitle("Preisliste Ping-Spiel")
c.setAuthor("Ping-Spiel")
c.showPage()
c.save()

print(DOCX_PATH)
print(PDF_PATH)
