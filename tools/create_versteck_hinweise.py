from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DOC_OUT = ROOT / "output" / "documents"
PDF_OUT = ROOT / "output" / "pdf"
DOC_OUT.mkdir(parents=True, exist_ok=True)
PDF_OUT.mkdir(parents=True, exist_ok=True)

DOCX_PATH = DOC_OUT / "Versteck_Hinweise_6x.docx"
PDF_PATH = PDF_OUT / "Versteck_Hinweise_6x.pdf"

TEXT = "Auf Bodenhöhe von Stein umgeben."
NAVY = "17324D"
BLUE = "2B6F93"
PALE = "F1F7FA"
BORDER = "AFC5D0"
INK = RGBColor(23, 50, 77)
MUTED = RGBColor(74, 101, 116)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "dashed")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, value=360):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size, bold=False, color=INK):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)
section.header_distance = Cm(0.5)
section.footer_distance = Cm(0.5)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for row in table.rows:
    row.height = Cm(8.9)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell in row.cells:
        cell.width = Cm(9.3)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, PALE)
        set_cell_borders(cell)
        set_cell_margins(cell, 420)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(13)
        label = p.add_run("HINWEIS")
        set_run(label, 10, True, MUTED)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        text = p.add_run(TEXT)
        set_run(text, 20, True, INK)

doc.core_properties.title = "Sechs Versteck-Hinweise"
doc.core_properties.subject = TEXT
doc.core_properties.author = "Ping-Spiel"
doc.save(DOCX_PATH)


def hc(value):
    return HexColor(f"#{value}")


c = canvas.Canvas(str(PDF_PATH), pagesize=A4)
page_w, page_h = A4
margin_x = 12 * mm
margin_y = 12 * mm
card_w = (page_w - 2 * margin_x) / 2
card_h = (page_h - 2 * margin_y) / 3

for row in range(3):
    for col in range(2):
        x = margin_x + col * card_w
        y = page_h - margin_y - (row + 1) * card_h

        c.setFillColor(hc(PALE))
        c.setStrokeColor(hc(BORDER))
        c.setLineWidth(0.8)
        c.setDash(4, 3)
        c.rect(x, y, card_w, card_h, fill=1, stroke=1)
        c.setDash()

        cx = x + card_w / 2
        cy = y + card_h / 2
        c.setFillColor(hc(BLUE))
        c.roundRect(cx - 13 * mm, cy + 14 * mm, 26 * mm, 8 * mm, 4 * mm, fill=1, stroke=0)
        c.setFillColor(hc("FFFFFF"))
        c.setFont("Helvetica-Bold", 9.5)
        c.drawCentredString(cx, cy + 16.5 * mm, "HINWEIS")

        c.setFillColor(hc(NAVY))
        c.setFont("Helvetica-Bold", 17)
        c.drawCentredString(cx, cy + 1 * mm, "Auf Bodenhöhe")
        c.drawCentredString(cx, cy - 8 * mm, "von Stein umgeben.")

c.setTitle("Sechs Versteck-Hinweise")
c.setAuthor("Ping-Spiel")
c.showPage()
c.save()

print(DOCX_PATH)
print(PDF_PATH)
